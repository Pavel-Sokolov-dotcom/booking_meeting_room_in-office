from rest_framework.exceptions import ValidationError
from rest_framework import generics
from .models import MeetingRoom, Booking
from .serializers import MeetingRoomSerializer, BookingSerializer
from rest_framework_simplejwt.views import TokenObtainPairView, TokenRefreshView
from rest_framework.permissions import IsAuthenticated
from django.utils import timezone
from datetime import timedelta
from rest_framework.response import Response
import logging
import os
from django.http import HttpResponse
from docx import Document
from dotenv import load_dotenv

load_dotenv()


logger = logging.getLogger(__name__)


class MeetingRoomList(generics.ListCreateAPIView):
    queryset = MeetingRoom.objects.all()
    serializer_class = MeetingRoomSerializer
    
    
class BookingList(generics.ListCreateAPIView):
    serializer_class = BookingSerializer
    permission_classes = [IsAuthenticated]
    
    def perform_create(self, serializer):
        room = serializer.validated_data['room']
        start_time = serializer.validated_data['start_time']
        end_time = serializer.validated_data['end_time']
        
        logger.info(f'Попытка создать бронирование: Комната {room}, Начало {start_time}, Конец {end_time}')

        if Booking.objects.filter(room=room, 
                                   start_time__lt=end_time, 
                                   end_time__gt=start_time).exists():
            logger.warning("Переговорка уже забронирована на это время.")
            raise ValidationError("Переговорка уже забронирована на это время.")

        serializer.save(user=self.request.user)
        logger.info(f'Создано бронирование: {serializer.validated_data}')
    
    def get_queryset(self):
        date_filter = self.request.query_params.get('date')
        if date_filter:
            return Booking.objects.filter(start_time__date=date_filter)
        return Booking.objects.all()
    
    def list(self, request, *args, **kwargs):
        rooms = MeetingRoom.objects.all()
        now = timezone.localtime()
        logging.info(f"Текущая дата: {now.date()}")
        bookings_today = Booking.objects.filter(start_time__date=now.date())
        logging.info(f"Сегодняшние бронирования: {bookings_today}")

        
        responce_date = []
        
        for room in rooms:
            is_available = not bookings_today.filter(room=room, start_time__lt=timezone.now(), end_time__gt=timezone.now()).exists()
            logging.info(f"Комната: {room.name}, Доступна: {is_available}")
            room_booking = [
                {
                    'room_id': room.id,
                    'user': booking.user.username,
                    'start_time': booking.start_time,
                    'end_time': booking.end_time,
                    'purpose': booking.purpose,
                } for booking in bookings_today.filter(room=room)
            ]
            logging.info(f"Бронирования для комнаты {room.name}: {room_booking}")
            responce_date.append({
                'room_id': room.id,
                'room': room.name,
                'is_available': is_available,
                'bookings': room_booking,
            })
            
        return Response(responce_date)


class BookingReport(generics.ListAPIView):
    serializer_class = BookingSerializer
    
    def get_queryset(self):
        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        if start_date and end_date:
            return Booking.objects.filter(start_time__range=[start_date, end_date])
        
        return Booking.objects.none()
    
    def get(self, request, *args, **kwargs):
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')
        bookings = self.get_queryset()

        # Создаем документ Word
        doc = Document()
        doc.add_heading('Отчет о бронированиях', level=1)
        doc.add_paragraph(f'Период: {start_date} - {end_date}')

        for booking in bookings:
            doc.add_paragraph(f'Комната: {booking.room.name}, '
                              f'Пользователь: {booking.user.username}, '
                              f'Начало: {booking.start_time}, '
                              f'Конец: {booking.end_time}, '
                              f'Цель: {booking.purpose}')

        # Задайте путь для сохранения документа
        directory_path = os.getenv('directory_path')
        logger.info(f"Проверка наличия директории: {directory_path}")
        if not os.path.exists(directory_path):
            logger.info(f"Создание папки: {directory_path}")
            os.makedirs(directory_path)  # Создание директории, если не существует

        file_name = f'booking_report_{timezone.now().strftime("%Y%m%d_%H%M%S")}.docx'
        file_path = os.path.join(directory_path, file_name)

        try:
            # Сохранение документа на сервере
            doc.save(file_path)
            logger.info(f"Документ сохранен: {file_path}")

            # Подготовка документа для скачивания
            with open(file_path, 'rb') as f:
                response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename={file_name}'
                return response
        except Exception as e:
            logger.error(f"Ошибка при сохранении документа: {e}")
            return HttpResponse(status=500)  # Возвращаем ошибку 500, если что-то пошло не так